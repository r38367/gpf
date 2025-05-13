

# History
#
# v1.0
#   03/12/23 add #1 open/edit jpg button  
#   04/12/23 add #5 clear button
#   05/12/23 add #4 open/edit word button
#   05/12/23 add #11 clear with double click on label
#
# v2.0
#   09/12/23 add #6 add file list in inout field
#   10/12/23 merge #22 missing 'sceren' in screen file name
#   10/12/23 fix #20 file not open, change file naming
#   
# v3.0
#   11/03/23 fix #14 add text to word file
# v3.1
#    27/03/25 fix #28 delete comment after screenshot
# v4.0
#   13/05/25 add #32 add button to open folder with files

nVer = '4.0'

# Import Module
from tkinter import *
import tkinter.messagebox as mb
import os, sys
import time

# ==============================================================================================================
# create a timestamp in form dd.mm.yy hh:mm:ss 
# ==============================================================================================================

import datetime

def timestamp(): 
    # Get the current date and time
    now = datetime.datetime.now()

    return now.strftime("%d.%m.%y %H:%M:%S")

# ==============================================================================================================
# take a screenshot 
# ==============================================================================================================


from PIL import ImageGrab

def takescreenshoot( file ):
    # Capture the entire screen
    screenshot = ImageGrab.grab()

    
    # Save the screenshot to a file
    screenshot.save( file )

    # Close the screenshot
    screenshot.close()

def makescreenshotname( t, text ):
    
    subfolder=".\\screens" # if run not from folder - os.path.dirname(sys.argv[0]) +
    if not os.path.exists(subfolder):
        os.mkdir(subfolder)
    
    if not text:
        text = "screen"

    return   subfolder + '\\' + t.strip().replace(':','').replace(' ','_') + "_" + text +".png"
 

# ==============================================================================================================
# working with word document 
# ==============================================================================================================

from docx import Document
from docx.shared import Mm
import re  

def get_text_width(document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000

def savetoword( file, text, jpg):
    
    try:
        # Load the existing document
        document = Document( file )
    except:
        # Create a new document if it does not exist
        document = Document()

    document.sections[0].left_margin = Mm(25)
    document.sections[0].right_margin = Mm(25)

    document.add_paragraph( text )
    document.paragraphs[-1].add_run().add_picture(jpg, Mm(get_text_width(document)) )
                        #width=Mm(document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin)) 

    try: 
            document.save(file)
    except:
            mb.showerror(title="Error", message="Can't save " + file + ".\nClose if open.")

def makewordname( input ):
    if input == "" :
        return ""
    else: 
        return input + '.docx'
    
# ==============================================================================================================
#  Call backs 
# ==============================================================================================================

test = 0
policy = 0
wordname = ""
jpg = ""


# ==============================================================================================================
#  Open Explorer 
# ==============================================================================================================

def _open_explorer():
    folder_path = os.getcwd()  # <-- Change this to your desired folder path
    if os.path.isdir(folder_path):
        os.startfile(folder_path)
    else:
        mb.showerror("Error", f"Folder does not exist:\n{folder_path}")

# ==============================================================================================================
#                   Input 
# pressed  
# ==============================================================================================================
   
def _input_pressed(event):
    global wordname 
    
    # hide listbox
    listbox.grid_remove()
    # get text from input 
    res = txt.get().strip()
    # make file name from input 
    wordname = makewordname( res )
    # set label to file name 
    lbl.configure(text = wordname)
    
    # define koeff from pix to char for entry widget
    # entry widget char=6 pixels 
    a_width = 6
    lbl_pix = lbl.winfo_reqwidth()
    #lbl_ch = len(lbl.cget("text") )
    #print( "pix",lbl_pix,"ch", lbl_ch, lbl_pix/lbl_ch  ) 
    txt_w = int(lbl.winfo_reqwidth()/a_width)

    # define width in char that matches lbl width in pixels 
    txt.config(width=max(txt_w, 40))

# ==============================================================================================================
#                   Take button 
# 
# pressed/released/changed
# ==============================================================================================================
  
_pressed = 0  # 0-not pressed, 1-short, 2-long

def _change_label(event):
    global _pressed
    if _pressed == 1:
        _pressed = 2
        btn1.configure( text='Edit')

     
def _take_button_pressed( event ):
    global _pressed
    global _long

    if _pressed == 0 :
        _pressed = 1
        _long = root.after(700,_change_label,event)
        

def _take_button_released( event ):

    global _pressed
    global _long

    #if short press
    if _pressed == 1:
        if _long:
            root.after_cancel(_long)    # cancel long event

    # take screenshot
    global wordname 
    global jpg
    global head

    _input_pressed( 0 )
    t = timestamp() 
    
    jpg = makescreenshotname( t, txt.get().strip() )
    takescreenshoot( jpg )


    # edit picture
    if _pressed == 2:
        os.system('"'+jpg+'"')
       
    _pressed = 0
    btn1.configure( text='Take')

    # save to word if needed 
    if not wordname == "":
        comment = head.get() 
        if( comment ):
            t += ' ' + comment.strip()
            head.delete(0,END)
        print( comment )
        
        savetoword( wordname, t, jpg )
        res = 'Saved to ' + wordname  
    else:
        res = 'Saved to ' + jpg 

    lbl.configure(text = res)


# ==============================================================================================================
#                   Open button
# 
#            pressed/released/changed
# ==============================================================================================================

_pressed_open = 0  # 0-not pressed, 1-short, 2-long

def _open_change_label(event):
    global _pressed_open
    if _pressed_open == 1:
        _pressed_open = 2
        btn2.configure( text='Word')

     
def _open_button_pressed( event ):
    global _pressed_open
    global _long

    _input_pressed( 0 )
    
    if _pressed_open == 0 :
        _pressed_open = 1
        _long = root.after(700,_open_change_label,event)
        

def _open_button_released( event ):

    global _pressed_open
    global _long

    #if short press
    if _pressed_open == 1:
        if _long:
            root.after_cancel(_long)    # cancel long event

    # open file
    global wordname 
    global jpg

    # open picture/word
    if _pressed_open == 2:
        
        if os.path.isfile(wordname):
            os.system('"'+ wordname +'"')
        else:
            # Show an information message box
            mb.showinfo(title="Message", message="No word file to show")

    else:
         
        if jpg == "":
            # Show an information message box
            mb.showinfo(title="Message", message="No image file to show")
        else:
            os.system('"'+jpg+'"')

    _pressed_open = 0
    btn2.configure( text='Open')
    
   

# ==============================================================================================================
#                  Clear button 
# pressed 
# ==============================================================================================================

def _clear_button_pressed(event=None):
    
    txt.delete(0, END)
    head.delete(0,END)
    lbl.configure(text = "")
    listbox.grid_remove()
    print( '-',head.get().strip() ,'-' ) 

# ==============================================================================================================
#                   Listbox - update
#
# update after input text in input field  
# ==============================================================================================================

def update_listbox (event=None):
    
    if event.keysym == "Return":
        # Move the focus to the listbox
        listbox.grid_remove()
        return

    if txt.get() == current_path:
        txt.delete(0, END)  # Clear the placeholder text
        txt.config(fg='black')  

    # Clear the listbox
    listbox.delete (0, END)
    # Get the input text
    input_text = txt.get ().strip()
    #check if string is empty
    if not input_text:
        return 
    # Initialize a counter for the number of matching files
    count = 0
    # Loop through the files in the current folder
    for file in os.listdir ("."):
        # Check if the file name contains the input text
        if all(file.upper().find(str) >= 0  for str in input_text.upper().split()):
            # Insert the file name into the listbox
            listbox.insert (END, file)
            # Increment the counter
            count += 1
    # Resize the listbox height according to the counter
    listbox.config (height=min(count,5)) 
    if count < 1:
        listbox.grid_remove()
    else:
        listbox.config(width=max( int(lbl.winfo_reqwidth()/6),40) )
        listbox.grid(column=0, row=2)

    
        
# ==============================================================================================================
#                   Listbox - return from listbox
# 
# Create a function to update the entry when press Return in listbox
# ==============================================================================================================
def update_entry (event):
    
    # Check if the event is triggered by pressing ENTER
    if event.keysym == "Return":
        # Get the selected file name from the listbox
        file_name = listbox.get (listbox.curselection ())
        # hide listbox
        listbox.grid_remove()
        # Set the input text without .docx extension
        txt.delete(0, END)
        txt.insert (0, file_name.rsplit('.',1)[0])
        # Update the listbox
        #update_listbox ()
        
        # Hide the toplevel window 
        # Changing position of cursor to end 
        txt.icursor(END)
        # set focus to entry
        txt.focus()
        #top.withdraw ()
        _input_pressed(None)
        

# ==============================================================================================================
#                   Listbox - show file name 
# 
# Show the active entry from listbox in the input  
# ==============================================================================================================
def show_file_name (event):
    #in case we do not have activ listbox selection keep text in input field
    index = listbox.curselection ()
    if index:
        # Get the selected file name from the listbox
        file_name = listbox.get (index)
        # Set the input text without .docx extension
        txt.delete(0, END)
        txt.insert (0, file_name.rsplit('.',1)[0])

# ==============================================================================================================
#                   Listbox - move focus 
# 
# Create a function to move the focus to the listbox
# ==============================================================================================================
def move_focus (event):
    # Check if the event is triggered by pressing DOWN
    if event.keysym == "Down":
        # Move the focus to the listbox
        listbox.focus ()
        # Select the first item in the listbox
        listbox.selection_set (0)
        
    if event.num == 1:
        index = listbox.nearest(event.y)
        value = listbox.get(index)
        listbox.selection_set (index)
        

# ==============================================================================================================
# Create GUI
# ==============================================================================================================

root = Tk()
 
#  windowSet  title 
root.title("Test DinSide - v." + nVer )

# Set geometry(widthxheight)
root.geometry('') 

# Set margins inside window
root['padx']=10
root['pady']=5

# make windows resizable 
root.resizable(False, False)
# handle resize event to re-allocate buttons
#root.bind("<Configure>", on_window_resize)

# Set always on top
root.attributes('-topmost',True)

# Set window bg color
#root.configure(bg='lightgray')

#
# widgets:
#      0                       1      2
# 0 [input text           ] [Take] [Open]
# 1 [label text           ]
#
# Create an entry widget
head = Entry(root) #,height=1) #,width=20)
head.grid(column=0,row=0, columnspan=3, padx=5, pady=5, sticky=EW)

#
#  create input widget
#

# Create an entry widget
txt = Entry(root, width=40, fg='gray')
# place top left
txt.grid(column =0, row =1, padx=5, pady=5, sticky=W)
# get current path
current_path = os.getcwd()
txt.insert(0,current_path)
# Bind to Return key
txt.bind('<Return>', _input_pressed)
# Bind the entry widget to the update function
txt.bind ("<KeyRelease>", update_listbox )
# Bind the entry widget to the move focus function
txt.bind ("<Down>", move_focus)
# Bind the root window to the show and hide functions
txt.bind ("<FocusIn>", update_listbox)

# Pack the entry widget
#txt.pack ()

#
# Create label 
#
lbl = Label(root, text = "", anchor="w", justify="left")
lbl.grid(column =0, row =2, sticky=W, padx=5)
#lbl.bind('<Double-1>', _clear_button_pressed) # Useless

#
# Create button Take
#
btn1 = Button(root, text = "Take" , width=6, fg = "black")
btn1.grid(column=1, row=1, padx=5)
btn1.bind('<Button-1>', _take_button_pressed)
btn1.bind('<ButtonRelease-1>', _take_button_released)

#
# Create button Open
#
btn2 = Button(root, text = "Open" , width=6, fg = "black")
btn2.grid(column=2, row=1, padx=5)
btn2.bind('<Button-1>', _open_button_pressed)
btn2.bind('<ButtonRelease-1>', _open_button_released)
#
# Create button Clear
#
btn3 = Button(root, text = "Clear" , width=6, fg = "black", command=_clear_button_pressed)
#btn3.grid(column=1, row=2, columnspan=2, padx=5, sticky=N)
btn3.grid(column=1, row=2, padx=5)

#
# Create button Clear
#
btn4 = Button(root, text = "Folder" , width=6, fg = "black", command=_open_explorer)
btn4.grid(column=2, row=2, padx=5)


#
# Create listbox
#
listbox = Listbox (root, width=40, height=5)
# Pack the listbox widget
listbox.grid(column=0, row=2, sticky=NW, padx=5, pady=5)
# Bind the listbox widget to the update function
listbox.bind ("<Return>", update_entry)
# Bind the listbox widget to the update function
listbox.bind ("<Button-1>", move_focus)
# Bind the listbox widget to the show file name function when item is selected 
listbox.bind ("<<ListboxSelect>>", show_file_name)
# Bind the listbox widget to quit when item is selected 
listbox.bind ("<FocusOut>", lambda e: listbox.grid_remove ())
            
# Update the listbox initially
listbox.grid_remove ()

# Start the main loop
root.mainloop ()
