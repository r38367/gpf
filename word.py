from docx import Document
from docx.shared import Mm
import re  

def get_text_width(document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000

def savetoword( file, text, jpg):
    
    try:
        # Load the existing document
        document = Document( file )
    except:
        # Create a new document if it does not exist
        document = Document()

    #document = Document('demo.docx')
    document.sections[0].left_margin = Mm(25)
    document.sections[0].right_margin = Mm(25)

    document.add_paragraph( text )
    document.paragraphs[-1].add_run().add_picture(jpg, Mm(get_text_width(document)) )
                        #width=Mm(document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin)) 

    document.save(file)


def makewordname( input ):
    if input == "" :
         return ""
    else: 
         
        # Find the first 4-digit number
        four_digit = re.findall(r'\b\d{4}\b', input)
        if not four_digit:
            testnum = ""
        else:
            testnum = four_digit[0]
        # Find the first 7-digit number
        seven_digit = re.findall(r'\b\d00\d{4}\b', input )
        if not seven_digit:
            policynum = ""
        else:
            policynum = seven_digit[0]
        
        # Concatenate the two numbers
        ret = testnum + policynum
    if ret == "":
        return ""
    else:
        return testnum +'_Anton_' + policynum + '.docx'
"""
from screen import takescreenshoot
from timestamp import timestamp 

text = timestamp()
jpg = text.replace(':','').replace(' ','_') + "_screen.jpg"
takescreenshoot( jpg )
savetoword( 'demo.docx', text, jpg )
"""